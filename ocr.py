import os
import re
import io
from pathlib import Path
import torch
from transformers import AutoTokenizer, AutoModel
from pdf2image import pdfinfo_from_path, convert_from_path
from docx import Document
from docx.shared import Pt
from tqdm import tqdm
from PIL import Image
from contextlib import redirect_stdout

# ===== CONFIG =====
PDF_PATH = r"test2.pdf"               # Input PDF path
OUTPUT_DOCX = r"output.docx"         # Word output file
OUTPUT_TXT = r"output.txt"           # Combined plain text file
MODEL_ID = r"./deepseek-ocr-model"   # Model folder path
# ==================

device = "cuda" if torch.cuda.is_available() else "cpu"
dtype = torch.bfloat16 if device == "cuda" else torch.float32
print(f"Device: {device} | Torch dtype: {dtype}")

# ===== Load DeepSeek Model =====
tokenizer = AutoTokenizer.from_pretrained(MODEL_ID, trust_remote_code=True)
model = AutoModel.from_pretrained(
    MODEL_ID,
    trust_remote_code=True,
    use_safetensors=True,
    attn_implementation="eager",
).to(dtype=dtype, device=device).eval()

# ===== Prepare Directories =====
info = pdfinfo_from_path(PDF_PATH)
page_count = info["Pages"]
output_dir = Path("ds_ocr_pages")
output_dir.mkdir(parents=True, exist_ok=True)

doc = Document()
prompt = "<image>\n<|grounding|>Convert the document page to markdown."
all_text = []

# ===== Extractor Function =====
def extract_text_from_stdout(raw: str) -> str:
    """
    Extract clean Hindi text that follows each <|ref|>text<|/ref|>...<|/det|> block.
    Removes markup and returns only readable OCR output.
    """
    pattern = re.compile(
        r"<\|ref\|>text<\|/ref\|>.*?<\|det\|>.*?\n(.*?)(?=<\|ref\||\Z)",
        flags=re.S,
    )
    matches = pattern.findall(raw)
    if not matches:
        # fallback: remove tags and return whatever is readable
        return re.sub(r"<.*?>", "", raw).strip()

    cleaned = "\n\n".join(
        re.sub(r"<.*?>", "", m).strip() for m in matches if m.strip()
    )
    return cleaned.strip()

# ===== OCR Loop =====
for p in tqdm(range(1, page_count + 1), desc="Extracting text"):
    # Convert PDF → Image
    images = convert_from_path(PDF_PATH, dpi=200, first_page=p, last_page=p)
    image = images[0].convert("RGB")
    img_path = output_dir / f"page_{p}.png"
    image.save(img_path)

    # Capture DeepSeek’s console output directly
    buffer = io.StringIO()
    with redirect_stdout(buffer):
        model.infer(
            tokenizer,
            prompt=prompt,
            image_file=str(img_path),
            output_path=str(output_dir),
            base_size=1024,
            image_size=640,
            crop_mode=True,
            save_results=False,  # disable file writes, capture stdout
            test_compress=False,
        )

    raw_output = buffer.getvalue()
    text = extract_text_from_stdout(raw_output)

    if not text:
        text = "[No text detected]"

    # ===== Save to DOCX =====
    doc.add_heading(f"Page {p}", level=2)
    para = doc.add_paragraph(text)
    if para.runs:
        run = para.runs[0]
        run.font.name = "Mangal"  # Devanāgarī-compatible font
        run.font.size = Pt(13)

    all_text.append(f"\n\n=== Page {p} ===\n{text}\n")

    if p % 5 == 0 or p == page_count:
        doc.save(OUTPUT_DOCX)
        print(f"Saved progress at page {p} ({len(text)} chars)")

# ===== Final Saves =====
doc.save(OUTPUT_DOCX)
with open(OUTPUT_TXT, "w", encoding="utf-8") as f:
    f.write("\n".join(all_text))

print("\nOCR completed successfully!")
print(f"Word file: {OUTPUT_DOCX}")
print(f"Text file: {OUTPUT_TXT}")
