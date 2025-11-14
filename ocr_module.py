import os
import io
import re
from pathlib import Path
import torch
from transformers import AutoTokenizer, AutoModel
from pdf2image import pdfinfo_from_path, convert_from_path
from docx import Document
from docx.shared import Pt
from tqdm import tqdm
from PIL import Image
from contextlib import redirect_stdout


# =========================================================
#  DeepSeek OCR Module
# =========================================================

def extract_text_from_stdout(raw: str) -> str:
    """
    Extract clean Hindi (or multilingual) text from DeepSeek model output.
    Removes markup and returns readable text only.
    """
    pattern = re.compile(
        r"<\|ref\|>text<\|/ref\|>.*?<\|det\|>.*?\n(.*?)(?=<\|ref\||\Z)",
        flags=re.S,
    )
    matches = pattern.findall(raw)
    if not matches:
        return re.sub(r"<.*?>", "", raw).strip()
    cleaned = "\n\n".join(
        re.sub(r"<.*?>", "", m).strip() for m in matches if m.strip()
    )
    return cleaned.strip()


def run_ocr(pdf_path: str, output_dir: str = "ocr_outputs") -> tuple[str, str]:
    """
    Perform OCR on the given PDF using DeepSeek OCR model.
    Returns paths to generated DOCX and TXT files.
    """
    # -------------------- Setup --------------------
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_docx = os.path.join(output_dir, f"{base_name}_output.docx")
    output_txt = os.path.join(output_dir, f"{base_name}_output.txt")

    # -------------------- Device --------------------
    device = "cuda" if torch.cuda.is_available() else "cpu"
    dtype = torch.bfloat16 if device == "cuda" else torch.float32
    print(f"[INFO] Device: {device} | Torch dtype: {dtype}")

    # -------------------- Load Model --------------------
    MODEL_ID = "deepseek-ocr-model"  # or HF model repo path
    tokenizer = AutoTokenizer.from_pretrained(
    MODEL_ID,
    trust_remote_code=True,
    use_fast=False
)

    model = AutoModel.from_pretrained(
        MODEL_ID,
        trust_remote_code=True,
        use_safetensors=True,
        attn_implementation="eager",
    ).to(dtype=dtype, device=device).eval()

    # -------------------- Prepare PDF --------------------
    info = pdfinfo_from_path(pdf_path)
    page_count = info["Pages"]
    pages_dir = Path(output_dir) / f"{base_name}_pages"
    pages_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    all_text = []
    prompt = "<image>\n<|grounding|>Convert the document page to markdown."

    # -------------------- OCR Loop --------------------
    for p in tqdm(range(1, page_count + 1), desc="Extracting text"):
        # Convert PDF → Image
        images = convert_from_path(pdf_path, dpi=200, first_page=p, last_page=p)
        image = images[0].convert("RGB")
        img_path = pages_dir / f"page_{p}.png"
        image.save(img_path)

        # Run inference and capture stdout
        buffer = io.StringIO()
        with redirect_stdout(buffer):
            model.infer(
                tokenizer,
                prompt=prompt,
                image_file=str(img_path),
                base_size=1024,
                image_size=640,
                crop_mode=True,
                save_results=False,
                output_path=str(pages_dir),
                test_compress=False,
            )

        raw_output = buffer.getvalue()
        text = extract_text_from_stdout(raw_output) or "[No text detected]"

        # Write to Word
        doc.add_heading(f"Page {p}", level=2)
        para = doc.add_paragraph(text)
        if para.runs:
            run = para.runs[0]
            run.font.name = "Mangal"  # Devanāgarī font
            run.font.size = Pt(13)

        # Collect text for TXT
        all_text.append(f"\n\n=== Page {p} ===\n{text}\n")

        # Save intermediate progress
        if p % 3 == 0 or p == page_count:
            doc.save(output_docx)

    # -------------------- Final Saves --------------------
    doc.save(output_docx)
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write("\n".join(all_text))

    print(f"[DONE] OCR completed → {output_docx}, {output_txt}")
    return output_docx, output_txt


# Optional: quick manual test (runs locally)
if __name__ == "__main__":
    test_pdf = "test2.pdf"
    if os.path.exists(test_pdf):
        run_ocr(test_pdf)
    else:
        print("Place a test.pdf in the same folder to test.")
